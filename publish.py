#!/usr/bin/env python3
"""
publish.py – Convert a Word .docx to a live article on timscudderproduct.com

Usage:
    python publish.py path/to/article.docx

The script will:
  1. Extract and convert the docx content to HTML
  2. Call Claude to generate SEO-optimised metadata
  3. Show you a preview and let you confirm / edit
  4. Write the article HTML, update the index, prev/next links, and sitemap
"""

import sys
import os
import re
import html
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.oxml.ns import qn


# ── Paths ────────────────────────────────────────────────────────────────────

REPO = Path(__file__).parent
ARTICLES_DIR = REPO / "articles"
INDEX_FILE = ARTICLES_DIR / "index.html"
SITEMAP_FILE = REPO / "sitemap.xml"
BASE_URL = "https://timscudderproduct.com"


# ── DOCX → HTML conversion ───────────────────────────────────────────────────

def _run_to_html(run):
    """Convert a single docx Run to an inline HTML fragment."""
    text = html.escape(run.text)
    if not text:
        return ""

    # Check for hyperlink (run's parent may be a hyperlink element)
    # Hyperlinks are handled at paragraph level; here just apply formatting.
    if run.bold and run.italic:
        text = f"<strong><em>{text}</em></strong>"
    elif run.bold:
        text = f"<strong>{text}</strong>"
    elif run.italic:
        text = f"<em>{text}</em>"
    return text


def _para_to_html(para):
    """Convert a docx Paragraph to an HTML string (may be empty string)."""
    style = para.style.name if para.style else "Normal"

    # Gather inline HTML, preserving hyperlinks
    inline_html = ""
    for child in para._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "hyperlink":
            # Extract the relationship ID to look up the URL later
            r_id = child.get(qn("r:id"))
            link_text = "".join(
                html.escape(node.text)
                for node in child.iter()
                if node.tag.endswith("}t") and node.text
            )
            # Bold/italic on runs inside the hyperlink
            runs_bold = any(
                node.tag.endswith("}b")
                for node in child.iter()
            )
            if r_id and para.part and hasattr(para.part, "rels"):
                try:
                    url = para.part.rels[r_id].target_ref
                    inner = f"<strong>{link_text}</strong>" if runs_bold else link_text
                    inline_html += f'<a href="{html.escape(url)}">{inner}</a>'
                except KeyError:
                    inline_html += link_text
            else:
                inline_html += link_text
        elif tag == "r":
            # Normal run
            from docx.text.run import Run as DocxRun
            run_obj = DocxRun(child, para)
            inline_html += _run_to_html(run_obj)
        # skip other elements (bookmarks, etc.)

    inline_html = inline_html.strip()

    if not inline_html:
        return ""

    # Map style names to HTML tags
    if style.startswith("Heading 1"):
        return f'<h1>{inline_html}</h1>'
    elif style.startswith("Heading 2"):
        return f'<h2 class="header-anchor-post">{inline_html}\n</h2>'
    elif style.startswith("Heading 3"):
        return f'<h3>{inline_html}</h3>'
    elif style.startswith("Heading 4"):
        return f'<h4>{inline_html}</h4>'
    elif style in ("Quote", "Intense Quote", "Block Text", "Block Quotation"):
        return f'<blockquote>\n<p>{inline_html}</p>\n</blockquote>'
    elif style in ("List Bullet", "List Bullet 2", "List Paragraph"):
        return f'<li><p>{inline_html}</p></li>'
    elif style in ("List Number", "List Number 2"):
        return f'<li><p>{inline_html}</p></li>'
    else:
        return f'<p>{inline_html}</p>'


def _is_list_style(style_name):
    return style_name and ("List Bullet" in style_name or "List Number" in style_name or "List Paragraph" in style_name)


def _is_numbered_style(style_name):
    return style_name and "List Number" in style_name


def docx_to_html(docx_path: Path) -> str:
    """Convert a .docx file to an HTML content fragment (no wrapping tags)."""
    doc = Document(str(docx_path))
    blocks = []
    paragraphs = doc.paragraphs

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        style = para.style.name if para.style else "Normal"

        if _is_list_style(style):
            # Collect consecutive list items
            numbered = _is_numbered_style(style)
            items = []
            while i < len(paragraphs):
                p = paragraphs[i]
                s = p.style.name if p.style else "Normal"
                if not _is_list_style(s):
                    break
                item_html = _para_to_html(p)
                if item_html:
                    items.append(item_html)
                i += 1
            if items:
                tag = "ol" if numbered else "ul"
                blocks.append(f"<{tag}>\n" + "\n".join(items) + f"\n</{tag}>")
            continue

        html_block = _para_to_html(para)
        if html_block:
            blocks.append(html_block)
        i += 1

    return "\n".join(blocks)


def docx_to_plain_text(docx_path: Path) -> str:
    """Extract plain text from a docx (for sending to Claude for analysis)."""
    doc = Document(str(docx_path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Claude metadata generation ────────────────────────────────────────────────

TAGS = ["Strategy", "Discovery", "Delivery", "AI", "Leadership", "Roadmaps", "Metrics"]

def generate_metadata(plain_text: str) -> dict:
    """Ask Claude to produce SEO-optimised metadata for the article."""
    client = anthropic.Anthropic()

    prompt = f"""You are an SEO expert helping to publish articles on a product management blog (timscudderproduct.com) by Tim Scudder, a product leader.

Analyse the following article and return metadata optimised for search engine discoverability and click-through rate.

Article content:
---
{plain_text[:8000]}
---

Return ONLY a JSON object with these exact keys (no markdown, no explanation):
{{
  "title": "The page/article title – compelling, clear, ≤60 chars where possible",
  "slug": "url-slug-using-hyphens-no-special-chars",
  "excerpt": "Meta description / article card excerpt – 1-2 sentences, ≤155 chars, hooks the reader",
  "tag": "ONE tag from this list: {TAGS}",
  "date": "{datetime.now().strftime('%B %Y')}"
}}

The title should be the human-readable article title (can be a phrase or sentence).
The slug should be derived from the title but can be shorter.
Pick the single most relevant tag from the list provided."""

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=512,
        messages=[{"role": "user", "content": prompt}]
    )

    import json
    raw = message.content[0].text.strip()
    # Strip markdown code fences if present
    raw = re.sub(r"^```[a-z]*\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)


# ── HTML template ─────────────────────────────────────────────────────────────

ARTICLE_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en-GB">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} – Tim Scudder</title>
    <link rel="icon" href="../favicon.ico?v=2" sizes="any">
    <link rel="icon" href="../favicon.svg?v=2" type="image/svg+xml">
    <link rel="icon" href="../favicon-32.png?v=2" type="image/png" sizes="32x32">
    <link rel="icon" href="../favicon-16.png?v=2" type="image/png" sizes="16x16">
    <link rel="apple-touch-icon" href="../apple-touch-icon.png?v=2" sizes="180x180">

    <!-- SEO & Social Sharing -->
    <meta name="description" content="{excerpt}">
    <meta name="author" content="Tim Scudder">
    <link rel="canonical" href="{base_url}/articles/{slug}.html">

    <!-- Open Graph -->
    <meta property="og:type" content="article">
    <meta property="og:title" content="{title} – Tim Scudder">
    <meta property="og:description" content="{excerpt}">
    <meta property="og:url" content="{base_url}/articles/{slug}.html">
    <meta property="og:image" content="{base_url}/og-image.png">
    <meta property="og:locale" content="en_GB">

    <!-- Twitter/X Card -->
    <meta name="twitter:card" content="summary_large_image">
    <meta name="twitter:title" content="{title} – Tim Scudder">
    <meta name="twitter:description" content="{excerpt}">
    <meta name="twitter:image" content="{base_url}/og-image.png">

    <link href="https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,400;0,500;0,600;1,400&family=DM+Serif+Display:ital@0;1&family=Lora:ital,wght@0,400;0,500;0,600;1,400&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="../styles.css">
    <!-- Google tag (gtag.js) -->
    <script async src="https://www.googletagmanager.com/gtag/js?id=G-XH1NVWMY9E"></script>
    <script>
      window.dataLayer = window.dataLayer || [];
      function gtag(){{dataLayer.push(arguments);}}
      gtag('js', new Date());
      gtag('config', 'G-XH1NVWMY9E');
    </script>
</head>
<body class="page-article">
    <div class="page-decoration"></div>
    <header class="site-header">
        <div class="site-header-bar">
            <a href="/" class="site-name">Tim Scudder, Thoughts on Product</a>
            <button class="nav-toggle" aria-expanded="false" aria-controls="site-nav" aria-label="Toggle navigation">
                <div class="nav-toggle-icon">
                    <span></span>
                    <span></span>
                    <span></span>
                </div>
            </button>
            <ul class="site-nav" id="site-nav">
                <li><a href="/articles">Articles</a></li>
                <li><a href="/coach.html">Product Sense Coach</a></li>
                <li><a href="/#contact">Contact</a></li>
                <li><a href="https://www.linkedin.com/in/timothyscudder/" target="_blank" rel="noopener">LinkedIn</a></li>
            </ul>
        </div>
    </header>
    <main>
    <article>
        <div class="article-meta">{date}<span class="article-tag">{tag}</span></div>
        <h1>{title}</h1>
        <div class="article-content">

{content}
        </div>
        <div class="article-contact-prompt">
            <p class="contact-prompt-heading">Thanks for reading!</p>
            <p class="contact-prompt-text">Want to chat about product? I'd love to hear from you.</p>
            <a href="/#contact" class="contact-prompt-button">Get in touch</a>
        </div>
        <div class="article-footer">
{footer_links}
</div>
    </article>
</main>
    <footer>
        <p>&copy; Tim Scudder</p>
    </footer>
    <script src="/nav.js"></script>
</body>
</html>
"""


# ── Index manipulation ────────────────────────────────────────────────────────

def get_first_article_slug(index_html: str) -> str | None:
    """Return the slug (filename without .html) of the first article card in the index."""
    match = re.search(r'<a href="([^"]+\.html)"[^>]*>\s*Read article\s*</a>', index_html)
    if match:
        return match.group(1).replace(".html", "")
    return None


def insert_article_card(index_html: str, slug: str, title: str, excerpt: str, tag: str) -> str:
    """Insert a new article card at the top of the articles-grid div."""
    card = f"""<div class="article-card">
    <span class="article-tag">{tag}</span>
    <h3><a href="{slug}.html">{html.escape(title)}</a></h3>
    <p class="article-excerpt">{html.escape(excerpt)}</p>
    <a href="{slug}.html" class="read-more">Read article</a>
</div>

"""
    # Insert after the opening tag of articles-grid
    return re.sub(
        r'(<div class="articles-grid">\s*\n)',
        r'\g<1>' + card,
        index_html,
        count=1
    )


# ── Prev/next link manipulation ───────────────────────────────────────────────

def add_next_link_to_article(article_path: Path, new_slug: str):
    """Add a 'Next article' link pointing to new_slug in an existing article."""
    content = article_path.read_text(encoding="utf-8")
    # If there's already a next link, don't add another
    if "Next article" in content:
        return
    # Find the article-footer div and add the next link
    content = re.sub(
        r'(<div class="article-footer">\s*\n)',
        rf'\g<1>  <a href="{new_slug}.html">Next article</a>\n',
        content,
        count=1
    )
    article_path.write_text(content, encoding="utf-8")


# ── Sitemap manipulation ──────────────────────────────────────────────────────

def add_to_sitemap(sitemap_html: str, slug: str) -> str:
    """Insert a new URL entry before the closing </urlset>."""
    entry = f"""
  <url>
    <loc>{BASE_URL}/articles/{slug}.html</loc>
    <changefreq>yearly</changefreq>
    <priority>0.7</priority>
  </url>

"""
    return sitemap_html.replace("</urlset>", entry + "</urlset>")


# ── Interactive confirmation ───────────────────────────────────────────────────

def prompt_edit(label: str, value: str) -> str:
    """Show current value; let user accept (Enter) or type a replacement."""
    print(f"  {label}: {value}")
    new = input(f"  → Keep this? Press Enter to accept, or type a replacement: ").strip()
    return new if new else value


def confirm_metadata(meta: dict) -> dict:
    print("\n── Claude's proposed metadata ──────────────────────────────")
    fields = ["title", "slug", "excerpt", "tag", "date"]
    for f in fields:
        meta[f] = prompt_edit(f.capitalize(), meta[f])
    print("────────────────────────────────────────────────────────────\n")
    return meta


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Usage: python publish.py path/to/article.docx")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: file not found: {docx_path}")
        sys.exit(1)

    print(f"\nReading {docx_path.name}...")
    plain_text = docx_to_plain_text(docx_path)
    content_html = docx_to_html(docx_path)

    print("Asking Claude to generate SEO metadata...")
    meta = generate_metadata(plain_text)

    meta = confirm_metadata(meta)

    slug = meta["slug"]
    article_file = ARTICLES_DIR / f"{slug}.html"

    if article_file.exists():
        overwrite = input(f"\nWarning: {article_file.name} already exists. Overwrite? (y/N): ").strip().lower()
        if overwrite != "y":
            print("Aborted.")
            sys.exit(0)

    # ── Read current state ────────────────────────────────────────────────────
    index_html = INDEX_FILE.read_text(encoding="utf-8")
    sitemap_xml = SITEMAP_FILE.read_text(encoding="utf-8")
    prev_first_slug = get_first_article_slug(index_html)

    # ── Build footer links ────────────────────────────────────────────────────
    if prev_first_slug:
        footer_links = f'  <a href="{prev_first_slug}.html">Previous article</a>'
    else:
        footer_links = ""

    # ── Write article HTML ────────────────────────────────────────────────────
    article_html = ARTICLE_TEMPLATE.format(
        title=html.escape(meta["title"]),
        slug=slug,
        excerpt=html.escape(meta["excerpt"]),
        tag=html.escape(meta["tag"]),
        date=meta["date"],
        base_url=BASE_URL,
        content=content_html,
        footer_links=footer_links,
    )
    article_file.write_text(article_html, encoding="utf-8")
    print(f"✓ Written {article_file.relative_to(REPO)}")

    # ── Update the previous first article to point back to this one ───────────
    if prev_first_slug:
        prev_article_path = ARTICLES_DIR / f"{prev_first_slug}.html"
        if prev_article_path.exists():
            add_next_link_to_article(prev_article_path, slug)
            print(f"✓ Added 'Next article' link to {prev_first_slug}.html")

    # ── Update articles index ─────────────────────────────────────────────────
    updated_index = insert_article_card(index_html, slug, meta["title"], meta["excerpt"], meta["tag"])
    INDEX_FILE.write_text(updated_index, encoding="utf-8")
    print(f"✓ Updated articles/index.html")

    # ── Update sitemap ────────────────────────────────────────────────────────
    updated_sitemap = add_to_sitemap(sitemap_xml, slug)
    SITEMAP_FILE.write_text(updated_sitemap, encoding="utf-8")
    print(f"✓ Updated sitemap.xml")

    print(f"\nDone! Article published at: {BASE_URL}/articles/{slug}.html")
    print(f"\nNext steps:")
    article_title = meta["title"]
    print(f"  git add -A && git commit -m 'Add article: {article_title}' && git push")


if __name__ == "__main__":
    main()
